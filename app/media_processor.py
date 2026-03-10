"""
Processamento de mídia para o agente: áudio (transcrição), imagem (descrição) e documentos (extração de texto).
"""
import io
import logging
from typing import Optional, Tuple

from openai import OpenAI

logger = logging.getLogger(__name__)

# Tipos que o agente consegue "ler"
MEDIA_AUDIO = "audio"
MEDIA_IMAGE = "image"
MEDIA_DOCUMENT = "document"


def _get_openai_client(api_key: Optional[str] = None) -> Optional[OpenAI]:
    if not api_key:
        return None
    return OpenAI(api_key=api_key)


def download_file(url: str, timeout: int = 30) -> Tuple[Optional[bytes], Optional[str]]:
    """
    Baixa o arquivo da URL. Retorna (bytes, content_type) ou (None, None) em erro.
    """
    try:
        import httpx
        r = httpx.get(url, timeout=timeout)
        r.raise_for_status()
        ct = (r.headers.get("content-type") or "").split(";")[0].strip().lower()
        return r.content, ct or None
    except Exception as e:
        logger.warning("download_file failed url=%s: %s", url[:80], e)
        return None, None


def transcribe_audio(data: bytes, api_key: str, content_type: Optional[str] = None) -> str:
    """
    Transcreve áudio com OpenAI Whisper. Retorna texto ou mensagem de fallback.
    """
    client = _get_openai_client(api_key)
    if not client:
        return "[Não foi possível transcrever o áudio: API key não configurada.]"
    try:
        # Whisper aceita mp3, mp4, mpeg, mpga, m4a, wav, webm
        ext = "mp3"
        if content_type:
            if "ogg" in content_type or "opus" in content_type:
                ext = "ogg"
            elif "wav" in content_type:
                ext = "wav"
            elif "m4a" in content_type or "x-m4a" in content_type:
                ext = "m4a"
            elif "webm" in content_type:
                ext = "webm"
        file_like = io.BytesIO(data)
        file_like.name = f"audio.{ext}"
        transcription = client.audio.transcriptions.create(
            model="whisper-1",
            file=file_like,
            response_format="text",
        )
        if isinstance(transcription, str):
            return transcription.strip() or "[Áudio sem fala detectada.]"
        # response_format="text" pode retornar objeto com .text
        out = (getattr(transcription, "text", None) or getattr(transcription, "content", "") or "")
        return (out if isinstance(out, str) else str(out)).strip() or "[Áudio sem fala detectada.]"
    except Exception as e:
        logger.exception("transcribe_audio failed: %s", e)
        return "[Não foi possível transcrever o áudio. Por favor, escreva sua mensagem em texto.]"


def describe_image(data: bytes, api_key: str, mime: Optional[str] = None) -> str:
    """
    Gera descrição da imagem com OpenAI Vision. Retorna texto para o agente.
    """
    client = _get_openai_client(api_key)
    if not client:
        return "[Não foi possível analisar a imagem: API key não configurada.]"
    try:
        import base64
        b64 = base64.standard_b64encode(data).decode("ascii")
        media_type = (mime or "image/jpeg").split(";")[0].strip()
        data_url = f"data:{media_type};base64,{b64}"
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {"url": data_url, "detail": "low"},
                        },
                        {
                            "type": "text",
                            "text": "Descreva esta imagem de forma objetiva e breve para um assistente que agenda entrevistas por WhatsApp. Inclua texto visível se houver (cartazes, documentos, telas).",
                        },
                    ],
                }
            ],
            max_tokens=500,
        )
        text = (response.choices[0].message.content or "").strip()
        return text or "[Imagem sem descrição disponível.]"
    except Exception as e:
        logger.exception("describe_image failed: %s", e)
        return "[Não foi possível analisar a imagem. Por favor, descreva em texto o que precisa.]"


def extract_document_text(data: bytes, mime_type: Optional[str], filename: Optional[str] = None) -> str:
    """
    Extrai texto de PDF, Excel, Word (e similares). Retorna texto ou mensagem de fallback.
    """
    mime = (mime_type or "").lower()
    fn = (filename or "").lower()

    try:
        if "pdf" in mime or fn.endswith(".pdf"):
            return _extract_pdf(data)
        if "spreadsheet" in mime or "excel" in mime or "vnd.ms-excel" in mime or fn.endswith((".xls", ".xlsx")):
            return _extract_excel(data)
        if "word" in mime or "msword" in mime or "wordprocessingml" in mime or fn.endswith((".doc", ".docx")):
            return _extract_word(data)
        # Tentar PDF por padrão para application/octet-stream com nome .pdf
        if fn.endswith(".pdf"):
            return _extract_pdf(data)
        return _extract_pdf(data) if "pdf" in mime else "[Formato de documento não suportado para leitura. Envie PDF, Excel ou Word.]"
    except Exception as e:
        logger.exception("extract_document_text failed: %s", e)
        return "[Não foi possível extrair o texto do arquivo. Por favor, copie as informações em uma mensagem de texto.]"


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        parts = []
        for i, page in enumerate(reader.pages):
            if i >= 20:
                parts.append("\n[... documento truncado ...]")
                break
            t = page.extract_text()
            if t:
                parts.append(t)
        return "\n".join(parts).strip() or "[PDF sem texto extraível.]"
    except ImportError:
        logger.warning("pypdf not installed")
        return "[Leitura de PDF não disponível. Instale pypdf.]"


def _extract_excel(data: bytes) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets[:5]:
            parts.append(f"--- Planilha: {sheet.title} ---")
            for row in sheet.iter_rows(max_row=200, values_only=True):
                line = "\t".join(str(c) if c is not None else "" for c in row)
                if line.strip():
                    parts.append(line)
        wb.close()
        return "\n".join(parts).strip() or "[Planilha sem conteúdo de texto.]"
    except ImportError:
        logger.warning("openpyxl not installed")
        return "[Leitura de Excel não disponível. Instale openpyxl.]"


def _extract_word(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables[:20]:
            for row in table.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(parts).strip() or "[Documento Word sem texto.]"
    except ImportError:
        logger.warning("python-docx not installed")
        return "[Leitura de Word não disponível. Instale python-docx.]"
